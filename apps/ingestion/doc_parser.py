"""
DOCX / DOC parser – extracts paragraphs and table content.
"""
from __future__ import annotations
from pathlib import Path
import io

from docx import Document
from apps.api.core.logging import get_logger

logger = get_logger(__name__)


def extract_docx(source: str | Path | bytes) -> str:
    """
    Return the full text of a DOCX file.

    Args:
        source: file path, Path, or raw bytes.

    Returns:
        Concatenated plain text of all paragraphs and table cells.
    """
    if isinstance(source, bytes):
        data = io.BytesIO(source)
    else:
        data = str(source)

    doc = Document(data)
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)
